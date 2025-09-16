"""
Document processing handler for DataRoom Intelligence Bot
Extracts content from PDFs (GPT-4o Direct only), Excel, Word, and other document types

PDF Processing Architecture: GPT-4o Direct Only
- Traditional methods (PyPDF2/pdfplumber/OCR) removed for superior structured extraction
- GPT-4o provides contextual financial data extraction vs raw text
- Simplified maintenance with single processing pipeline
"""

import os
# import pandas as pd  # COMENTAR TEMPORALMENTE
from typing import Dict, List, Optional, Any
from pathlib import Path
# import PyPDF2  # REMOVED: GPT-4o Direct only architecture
import docx
from utils.logger import get_logger

logger = get_logger(__name__)

# GPT-4o Direct PDF Processing (Only Method)
try:
    from handlers.gpt4o_pdf_processor import GPT4oDirectProcessor
    from config.settings import config
    gpt4o_available = config.openai_configured
    logger.info("✅ GPT-4o Direct processor available")
except ImportError as e:
    logger.warning(f"⚠️ GPT-4o Direct not available: {e}")
    gpt4o_available = False
except Exception as e:
    logger.error(f"❌ GPT-4o Direct initialization failed: {e}")
    gpt4o_available = False

class DocumentProcessor:
    """Processes various document types and extracts structured content

    PDF Processing: GPT-4o Direct only (fallback methods removed)
    - Architectural decision: GPT-4o provides superior structured data extraction
    - Evidence: Stamp PDF analysis showed contextual financial extraction vs raw numbers
    - Benefits: Maintenance simplification, structured output, slide references
    """

    def __init__(self):
        self.supported_extensions = {
            '.pdf': self._process_pdf,
            # '.xlsx': self._process_excel,    # COMENTAR TEMPORALMENTE
            # '.xls': self._process_excel,     # COMENTAR TEMPORALMENTE
            '.docx': self._process_word,
            '.doc': self._process_word,
            '.txt': self._process_text,
            # '.csv': self._process_csv        # COMENTAR TEMPORALMENTE
        }
        
        # Initialize GPT-4o Direct processing (single method)
        self.gpt4o_processor = None

        # Try to initialize GPT-4o processor
        try:
            if 'GPT4oDirectProcessor' in globals() and 'config' in globals():
                self.gpt4o_processor = GPT4oDirectProcessor(config.OPENAI_API_KEY)
                logger.info("✅ GPT-4o Direct PDF processor initialized (only method)")
            else:
                logger.info("ℹ️ GPT-4o Direct processor not available")
        except Exception as e:
            logger.warning(f"⚠️ Failed to initialize GPT-4o processor: {e}")
            self.gpt4o_processor = None

    def process_document(self, file_path: str, file_name: str, mime_type: str) -> Dict[str, Any]:
        """Process a document and extract its content"""
        try:
            logger.info(f"📄 Processing document: {file_name}")

            # Get file extension
            file_ext = Path(file_name).suffix.lower()

            if file_ext not in self.supported_extensions:
                logger.warning(f"⚠️ Unsupported file extension: {file_ext}")
                return {
                    'name': file_name,
                    'type': 'unsupported',
                    'content': '',
                    'metadata': {'error': f'Unsupported file type: {file_ext}'}
                }

            # Process based on file type
            processor_func = self.supported_extensions[file_ext]
            content_data = processor_func(file_path, file_name)

            logger.info(f"✅ Successfully processed: {file_name}")
            return content_data

        except Exception as e:
            logger.error(f"❌ Failed to process {file_name}: {e}")
            return {
                'name': file_name,
                'type': 'error',
                'content': '',
                'metadata': {'error': str(e)}
            }

    def _process_pdf(self, file_path: str, file_name: str) -> Dict[str, Any]:
        """Extract text content from PDF files using GPT-4o Direct processing only"""
        try:
            # GPT-4o Direct PDF Processing (primary method)
            if self.gpt4o_processor:
                logger.info(f"📄 Processing PDF with GPT-4o Direct: {file_name}")
                try:
                    gpt4o_result = self.gpt4o_processor.process_pdf_document(file_path, file_name)

                    # Check if GPT-4o processing succeeded
                    if (gpt4o_result and
                        gpt4o_result.get('content') and
                        len(gpt4o_result['content']) > 50 and  # Lower threshold for acceptance
                        not gpt4o_result.get('metadata', {}).get('fallback_required', False)):

                        logger.info(f"✅ GPT-4o Direct processing successful: {len(gpt4o_result['content'])} chars")
                        return gpt4o_result
                    else:
                        logger.warning(f"⚠️ GPT-4o processing failed or returned insufficient content")

                        # Return structured empty result instead of fallback
                        return {
                            'name': file_name,
                            'type': 'pdf',
                            'content': '',
                            'structured_data': None,
                            'metadata': {
                                'extraction_method': 'gpt4o_failed',
                                'file_size_bytes': os.path.getsize(file_path),
                                'error': 'GPT-4o processing failed to extract sufficient content',
                                'has_content': False,
                                'processing_attempted': True,
                                'gpt4o_error': gpt4o_result.get('metadata', {}).get('error', 'Unknown error')
                            }
                        }

                except Exception as e:
                    logger.error(f"❌ GPT-4o Direct processing failed: {e}")
                    return {
                        'name': file_name,
                        'type': 'pdf',
                        'content': '',
                        'structured_data': None,
                        'metadata': {
                            'extraction_method': 'gpt4o_failed',
                            'file_size_bytes': os.path.getsize(file_path),
                            'error': f'GPT-4o processing exception: {str(e)}',
                            'has_content': False,
                            'processing_attempted': True
                        }
                    }
            else:
                # No GPT-4o processor available
                logger.error(f"❌ GPT-4o processor not available for {file_name}")
                return {
                    'name': file_name,
                    'type': 'pdf',
                    'content': '',
                    'structured_data': None,
                    'metadata': {
                        'extraction_method': 'unavailable',
                        'file_size_bytes': os.path.getsize(file_path),
                        'error': 'GPT-4o processor not initialized',
                        'has_content': False,
                        'processing_attempted': False
                    }
                }

        except Exception as e:
            logger.error(f"❌ PDF processing failed for {file_name}: {e}")
            return {
                'name': file_name,
                'type': 'pdf',
                'content': '',
                'structured_data': None,
                'metadata': {
                    'extraction_method': 'system_error',
                    'error': str(e),
                    'has_content': False,
                    'debug_info': [f"Fatal error: {str(e)}"]
                }
            }

    # REMOVED: Traditional PDF extraction methods (_try_pypdf2_extraction, _try_pdfplumber_extraction, _try_ocr_extraction)
    # Architecture simplified to GPT-4o Direct only per architectural decision
    # Reasoning: GPT-4o provides superior structured data extraction vs raw text extraction
    # Evidence: Stamp PDF analysis showed GPT-4o extracts financial data with context vs traditional methods' raw numbers

    def _legacy_extraction_removed(self, file_path: str, file_name: str) -> Dict[str, Any]:
        """Legacy extraction methods removed - GPT-4o Direct only architecture"""
        logger.warning(f"⚠️ Legacy extraction method called for {file_name} - this should not happen")
        return {
            'name': file_name,
            'type': 'pdf',
            'content': '',
            'metadata': {
                'extraction_method': 'legacy_removed',
                'error': 'Traditional extraction methods removed in favor of GPT-4o Direct',
                'has_content': False
            }
        }

    def _process_word(self, file_path: str, file_name: str) -> Dict[str, Any]:
        """Extract text content from Word documents"""
        try:
            doc = docx.Document(file_path)

            content = f"Word Document: {file_name}\n"
            content += "=" * 50 + "\n\n"

            paragraph_count = 0
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content += paragraph.text + "\n\n"
                    paragraph_count += 1

            # Extract tables if any
            table_count = len(doc.tables)
            if table_count > 0:
                content += "\n--- TABLES ---\n\n"
                for i, table in enumerate(doc.tables):
                    content += f"Table {i + 1}:\n"
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells)
                        content += row_text + "\n"
                    content += "\n"

            return {
                'name': file_name,
                'type': 'word',
                'content': content.strip(),
                'metadata': {
                    'paragraphs': paragraph_count,
                    'tables': table_count,
                    'content_length': len(content),
                    'has_content': bool(content.strip())
                }
            }

        except Exception as e:
            logger.error(f"❌ Word processing failed: {e}")
            raise

    def _process_text(self, file_path: str, file_name: str) -> Dict[str, Any]:
        """Process plain text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()

            return {
                'name': file_name,
                'type': 'text',
                'content': content,
                'metadata': {
                    'content_length': len(content),
                    'line_count': len(content.split('\n')),
                    'has_content': bool(content.strip())
                }
            }

        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as file:
                content = file.read()

            return {
                'name': file_name,
                'type': 'text',
                'content': content,
                'metadata': {
                    'content_length': len(content),
                    'line_count': len(content.split('\n')),
                    'has_content': bool(content.strip()),
                    'encoding': 'latin-1'
                }
            }

    def process_dataroom_documents(self, downloaded_files: List[Dict]) -> List[Dict[str, Any]]:
        """Process all documents in a data room"""
        processed_documents = []

        logger.info(f"🔄 Processing {len(downloaded_files)} documents...")

        for file_info in downloaded_files:
            try:
                processed_doc = self.process_document(
                    file_info['path'],
                    file_info['name'],
                    file_info['mime_type']
                )
                processed_documents.append(processed_doc)

            except Exception as e:
                logger.error(f"❌ Failed to process {file_info['name']}: {e}")
                # Add error document to maintain document list
                processed_documents.append({
                    'name': file_info['name'],
                    'type': 'error',
                    'content': '',
                    'metadata': {'error': str(e)}
                })

        logger.info(f"✅ Processed {len(processed_documents)} documents")
        return processed_documents

    def get_content_summary(self, processed_documents: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Generate a summary of all processed documents"""
        summary = {
            'total_documents': len(processed_documents),
            'document_types': {},
            'total_content_length': 0,
            'successful_processing': 0,
            'failed_processing': 0,
            'document_list': []
        }

        for doc in processed_documents:
            doc_type = doc.get('type', 'unknown')

            # Count by type
            if doc_type not in summary['document_types']:
                summary['document_types'][doc_type] = 0
            summary['document_types'][doc_type] += 1

            # Track processing success
            if doc_type == 'error':
                summary['failed_processing'] += 1
            else:
                summary['successful_processing'] += 1
                summary['total_content_length'] += len(doc.get('content', ''))

            # Document list for reference
            summary['document_list'].append({
                'name': doc['name'],
                'type': doc_type,
                'has_content': bool(doc.get('content', '').strip()),
                'content_length': len(doc.get('content', ''))
            })

        return summary
